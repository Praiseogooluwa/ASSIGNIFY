from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from supabase import create_client, Client
from dotenv import load_dotenv
import os, io, zipfile, re, secrets, hashlib
from datetime import datetime, timezone, timedelta
from typing import Optional
from collections import defaultdict
import httpx
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
import jwt
import pdfplumber
from docx import Document as DocxDocument
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

load_dotenv()

# ─── Rate limiter ─────────────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)

app = FastAPI(title="Assignify API")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ─── CORS — locked to your domain ─────────────────────────────────────────────
ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS",
    "https://assignify.com.ng,https://www.assignify.com.ng,http://localhost:8080,http://localhost:5173"
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Security — all secrets from environment variables ────────────────────────
SECRET_KEY = os.getenv("SECRET_KEY", secrets.token_hex(32))
ALGORITHM = "HS256"

# ─── File upload limit: 10MB ──────────────────────────────────────────────────
MAX_UPLOAD_SIZE = 10 * 1024 * 1024

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

# ─── ADMIN CREDENTIALS ────────────────────────────────────────────────────────
ADMIN_EMAIL = os.getenv("ADMIN_EMAIL", "admin@assignify.app")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "changeme_set_in_render")
ADMIN_SECRET = os.getenv("ADMIN_SECRET", "admin_super_secret_key_change_this")

supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

security = HTTPBearer()


# ─── Departments ──────────────────────────────────────────────────────────────

DEPARTMENTS = [
    "Biology Education",
    "Chemistry Education",
    "Computer Science Education",
    "Educational Technology",
    "Mathematics Education",
    "Physics Education",
    "English Education",
    "Yoruba Education",
    "Social Studies Education",
    "Christian Religious Studies Education",
    "Islamic Studies Education",
    "History Education",
    "Political Science Education",
    "Educational Management",
    "Guidance & Counselling",
    "Educational Psychology",
    "Early Childhood Education",
    "Physical & Health Education",
    "Sports Science",
]


# ─── Auth helpers ─────────────────────────────────────────────────────────────

def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """
    Verify token — supports both:
    1. Supabase JWT (normal lecturer login)
    2. Custom JWT (admin impersonation via create_access_token)
    """
    token = credentials.credentials
    
    # First try our custom JWT (impersonation tokens are signed with SECRET_KEY)
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = payload.get("sub")
        if user_id:
            # Return a minimal user-like object for impersonated sessions
            class ImpersonatedUser:
                def __init__(self, uid, email):
                    self.id = uid
                    self.email = email
                    self.user_metadata = {}
            return ImpersonatedUser(user_id, payload.get("email", ""))
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token has expired")
    except jwt.InvalidTokenError:
        pass  # Not our custom JWT — try Supabase next

    # Fall back to Supabase JWT verification
    try:
        user = supabase.auth.get_user(token)
        if not user or not user.user:
            raise HTTPException(status_code=401, detail="Invalid or expired token")
        return user.user
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    

def create_access_token(data: dict, expires_delta: timedelta = timedelta(hours=12)):
    to_encode = data.copy()
    expire = datetime.utcnow() + expires_delta
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def verify_admin_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    """Verify our custom admin token (simple HMAC-signed string)."""
    token = credentials.credentials
    # Admin token format: "admin:<timestamp>:<signature>"
    try:
        parts = token.split(":")
        if len(parts) != 3 or parts[0] != "admin":
            raise HTTPException(status_code=403, detail="Not an admin token")
        timestamp = int(parts[1])
        signature = parts[2]
        # Check token is not older than 8 hours
        now_ts = int(datetime.now(timezone.utc).timestamp())
        if now_ts - timestamp > 28800:  # 8 hours
            raise HTTPException(status_code=401, detail="Admin session expired. Please log in again.")
        # Verify signature
        expected = hashlib.sha256(f"admin:{timestamp}:{ADMIN_SECRET}".encode()).hexdigest()[:32]
        if not secrets.compare_digest(signature, expected):
            raise HTTPException(status_code=403, detail="Invalid admin token")
        return True
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=403, detail="Invalid admin token")


# ─── Auth endpoints ───────────────────────────────────────────────────────────

@app.post("/auth/register")
@limiter.limit("5/minute")
async def register(
    request: Request,
    full_name: str = Form(...),
    email: str = Form(...),
    password: str = Form(...),
):
    """
    Register a new lecturer.
    - Returns 409 if email already exists
    - Supabase sends OTP email automatically
    """
    try:
        response = supabase.auth.sign_up({
            "email": email,
            "password": password,
            "options": {
                "data": {"full_name": full_name},
                "email_redirect_to": None,
            }
        })

        # Supabase silently "succeeds" for existing emails but returns
        # an empty identities list — catch this and return 409
        if response.user is None:
            raise HTTPException(status_code=400, detail="Registration failed. Please try again.")
        
        identities = getattr(response.user, "identities", None)
        if identities is not None and len(identities) == 0:
            raise HTTPException(
                status_code=409,
                detail="An account with this email already exists. Please sign in instead."
            )

        return {"message": "Verification code sent to your email"}
    except HTTPException:
        raise
    except Exception as e:
        detail = str(e).lower()
        if "already registered" in detail or "already been registered" in detail or "user already registered" in detail:
            raise HTTPException(
                status_code=409,
                detail="An account with this email already exists. Please sign in instead."
            )
        raise HTTPException(status_code=400, detail="Registration failed. Please try again.")


@app.post("/auth/verify")
async def verify_otp(
    email: str = Form(...),
    code: str = Form(...),
):
    """Verify the 6-digit OTP sent to lecturer's email after registration."""
    try:
        response = supabase.auth.verify_otp({
            "email": email,
            "token": code,
            "type": "signup",
        })
        if not response.session:
            raise HTTPException(status_code=400, detail="Invalid or expired code. Please try again.")
        return {
            "token": response.session.access_token,
            "user": {
                "id": response.user.id,
                "email": response.user.email,
                "full_name": response.user.user_metadata.get("full_name", ""),
            }
        }
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid or expired code. Please try again.")


@app.post("/auth/resend-code")
async def resend_code(email: str = Form(...)):
    """Resend OTP verification code."""
    try:
        supabase.auth.resend({"type": "signup", "email": email})
        return {"message": "New code sent"}
    except Exception:
        raise HTTPException(status_code=400, detail="Failed to resend code")


@app.post("/auth/login")
@limiter.limit("10/minute")
async def login(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
):
    """
    Login and return JWT token.
    - Returns specific error if email not confirmed yet
    - Returns generic error for wrong password (security best practice)
    """
    try:
        response = supabase.auth.sign_in_with_password({
            "email": email,
            "password": password,
        })
        if not response.session:
            raise HTTPException(status_code=401, detail="Invalid credentials")
        full_name = response.user.user_metadata.get("full_name", "") if response.user.user_metadata else ""
        return {
            "token": response.session.access_token,
            "user": {
                "id": response.user.id,
                "email": response.user.email,
                "full_name": full_name,
                "display_name": full_name,
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        detail = str(e).lower()
        if "email not confirmed" in detail:
            raise HTTPException(
                status_code=401,
                detail="Please verify your email first. Check your inbox for a 6-digit code."
            )
        raise HTTPException(status_code=401, detail="Invalid email or password")


@app.post("/auth/forgot-password")
async def forgot_password(email: str = Form(...)):
    try:
        supabase.auth.reset_password_for_email(email)
        return {"message": "If this email exists, a reset link has been sent"}
    except Exception as e:
        print("ERROR:", str(e))
        raise HTTPException(
            status_code=500,
            detail="Something went wrong. Please try again."
        )


# ─── Departments ──────────────────────────────────────────────────────────────

@app.get("/departments")
async def get_departments():
    return DEPARTMENTS


# ─── Assignments ─────────────────────────────────────────────────────────────

@app.get("/assignments")
async def get_assignments(user=Depends(get_current_user)):
    result = supabase.table("assignments")\
        .select("*")\
        .eq("lecturer_id", str(user.id))\
        .order("created_at", desc=True)\
        .execute()
    assignments = result.data

    for a in assignments:
        count_result = supabase.table("submissions")\
            .select("id", count="exact")\
            .eq("assignment_id", a["id"])\
            .execute()
        a["submission_count"] = count_result.count or 0

    return assignments


@app.post("/assignments")
async def create_assignment(
    course_name: str = Form(...),
    title: str = Form(...),
    submission_type: str = Form(...),
    deadline: str = Form(...),
    number_of_groups: Optional[int] = Form(None),
    instructions: Optional[str] = Form(None),
    target_level: Optional[str] = Form(None),  # e.g. "100L", "200L", "300L", "400L"
    user=Depends(get_current_user)
):
    data = {
        "course_name": course_name,
        "title": title,
        "submission_type": submission_type,
        "deadline": deadline,
        "number_of_groups": number_of_groups,
        "instructions": instructions,
        "target_level": target_level.strip() if target_level and target_level.strip() else None,
        "is_closed": False,
        "lecturer_id": str(user.id),
    }
    result = supabase.table("assignments").insert(data).execute()
    return result.data[0]


@app.patch("/assignments/{assignment_id}/close")
async def close_assignment(
    assignment_id: str,
    user=Depends(get_current_user)
):
    """Only the lecturer who owns this assignment can close it."""
    existing = supabase.table("assignments")\
        .select("lecturer_id")\
        .eq("id", assignment_id)\
        .single().execute()
    if not existing.data:
        raise HTTPException(status_code=404, detail="Assignment not found")
    if existing.data["lecturer_id"] != str(user.id):
        raise HTTPException(status_code=403, detail="You don't have permission to close this assignment")

    result = supabase.table("assignments")\
        .update({"is_closed": True})\
        .eq("id", assignment_id)\
        .execute()
    return result.data[0]


@app.get("/assignments/{assignment_id}")
async def get_assignment(assignment_id: str):
    """PUBLIC endpoint — students need this to load assignment details."""
    result = supabase.table("assignments").select("*").eq("id", assignment_id).single().execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Assignment not found")
    return result.data


@app.delete("/assignments/{assignment_id}")
async def delete_assignment(assignment_id: str, user=Depends(get_current_user)):
    """Only the owner can delete their assignment."""
    existing = supabase.table("assignments")\
        .select("lecturer_id")\
        .eq("id", assignment_id)\
        .single().execute()
    if not existing.data:
        raise HTTPException(status_code=404, detail="Assignment not found")
    if existing.data["lecturer_id"] != str(user.id):
        raise HTTPException(status_code=403, detail="You don't have permission to delete this assignment")

    supabase.table("submissions").delete().eq("assignment_id", assignment_id).execute()
    supabase.table("assignments").delete().eq("id", assignment_id).execute()
    return {"message": "Assignment deleted"}


# ─── Submissions ─────────────────────────────────────────────────────────────

@app.post("/submissions/{assignment_id}")
async def submit_assignment(
    assignment_id: str,
    full_name: str = Form(...),
    matric_number: str = Form(...),
    department: str = Form(...),
    group_number: Optional[int] = Form(None),
    file: UploadFile = File(...)
):
    # 1. Load assignment
    assignment = supabase.table("assignments").select("*").eq("id", assignment_id).single().execute()
    if not assignment.data:
        raise HTTPException(status_code=404, detail="Assignment not found")
    a = assignment.data

    # 2. Check if manually closed
    if a.get("is_closed"):
        raise HTTPException(status_code=400, detail="This assignment has been closed by the lecturer")

    # 3. Check deadline
    deadline = datetime.fromisoformat(a["deadline"].replace("Z", "+00:00"))
    now = datetime.now(timezone.utc)
    is_late = now > deadline
    if is_late:
        raise HTTPException(status_code=400, detail="Submission deadline has passed")

    # 4. Check duplicate matric
    existing = supabase.table("submissions")\
        .select("id")\
        .eq("assignment_id", assignment_id)\
        .eq("matric_number", matric_number.upper())\
        .execute()
    if existing.data:
        raise HTTPException(
            status_code=400,
            detail="This matric number has already submitted for this assignment"
        )

    # 5. Validate file type
    allowed_types = [
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ]
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only PDF, DOC, and DOCX files are accepted")

    # 6. Upload file to Supabase Storage
    content = await file.read()
    if len(content) > MAX_UPLOAD_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum size is 10MB. Your file is {round(len(content)/(1024*1024), 1)}MB."
        )
    safe_name = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename or "file")
    file_path = f"{assignment_id}/{matric_number.upper()}_{safe_name}"

    supabase.storage.from_("submissions").upload(
        file_path,
        content,
        {"content-type": file.content_type, "upsert": "true"}
    )

    file_url = supabase.storage.from_("submissions").get_public_url(file_path)

    # 7. Save submission record
    result = supabase.table("submissions").insert({
        "assignment_id": assignment_id,
        "full_name": full_name,
        "matric_number": matric_number.upper(),
        "department": department,
        "group_number": group_number,
        "file_url": file_url,
        "is_late": is_late,
        "submitted_at": datetime.now(timezone.utc).isoformat(),
    }).execute()

    return result.data[0]


@app.get("/submissions/{assignment_id}")
async def get_submissions(
    assignment_id: str,
    department: Optional[str] = None,
    group_number: Optional[int] = None,
    user=Depends(get_current_user)
):
    query = supabase.table("submissions").select("*").eq("assignment_id", assignment_id)
    if department:
        query = query.eq("department", department)
    if group_number:
        query = query.eq("group_number", group_number)
    result = query.order("submitted_at", desc=True).execute()
    return result.data


@app.patch("/submissions/{submission_id}/score")
async def update_score(
    submission_id: str,
    score: float = Form(...),
    user=Depends(get_current_user)
):
    result = supabase.table("submissions")\
        .update({"score": score})\
        .eq("id", submission_id)\
        .execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Submission not found")
    return result.data[0]


@app.get("/download/{assignment_id}/zip")
async def download_zip(
    assignment_id: str,
    mode: str = "all",
    filter_value: Optional[str] = None,
    user=Depends(get_current_user)
):
    """
    Download submissions as ZIP with proper folder structure:

    mode=all      → fetches ALL submissions
                    ZIP: By_Department / {Dept} / {Group} / {course}_{title}_{matric}_{name}.pdf

    mode=department (filter_value=dept name) → fetches only that department
                    ZIP: By_Group / {Group} / {course}_{title}_{matric}_{name}.pdf

    mode=group (filter_value=group number) → fetches only that group
                    ZIP: By_Department / {Dept} / {course}_{title}_{matric}_{name}.pdf
    """
    query = supabase.table("submissions").select("*").eq("assignment_id", assignment_id)
    if mode == "department" and filter_value:
        query = query.eq("department", filter_value)
    elif mode == "group" and filter_value:
        try:
            query = query.eq("group_number", int(filter_value))
        except ValueError:
            pass
    subs = query.execute().data

    if not subs:
        raise HTTPException(status_code=404, detail="No submissions found")

    # Get assignment info for file naming
    assign_info = supabase.table("assignments").select("course_name, title")        .eq("id", assignment_id).single().execute()
    a_data = assign_info.data or {}
    course = re.sub(r"[^a-zA-Z0-9]", "_", a_data.get("course_name", "course"))
    title_safe = re.sub(r"[^a-zA-Z0-9]", "_", a_data.get("title", "assignment"))

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        async with httpx.AsyncClient() as client:
            for sub in subs:
                try:
                    resp = await client.get(sub["file_url"], timeout=30)
                    if resp.status_code != 200:
                        continue

                    ext = sub["file_url"].split(".")[-1].split("?")[0].lower()
                    matric = re.sub(r"[^a-zA-Z0-9]", "_", sub.get("matric_number", "unknown"))
                    name = re.sub(r"[^a-zA-Z0-9]", "_", sub.get("full_name", "unknown"))
                    dept = re.sub(r"[^a-zA-Z0-9]", "_", sub.get("department", "Unknown_Department"))
                    grp = sub.get("group_number")
                    group_folder = f"Group_{grp}" if grp else "Individual"

                    # File is named: CourseName_AssignmentTitle_MatricNumber_FullName.ext
                    filename = f"{course}_{title_safe}_{matric}_{name}.{ext}"

                    if mode == "all":
                        # ALL: top level = Groups, inside each group = all files from any dept
                        zip_path = f"{group_folder}/{filename}"

                    elif mode == "department":
                        # One dept selected: just flat files, no subfolders
                        zip_path = filename

                    elif mode == "group":
                        # One group selected: just flat files, no subfolders
                        zip_path = filename

                    else:
                        zip_path = filename

                    zf.writestr(zip_path, resp.content)
                except Exception:
                    continue

    zip_buf.seek(0)
    zip_name = f"{course}_{title_safe}_submissions.zip"
    return StreamingResponse(
        zip_buf,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={zip_name}"}
    )


# ─── Shared class list parser ─────────────────────────────────────────────────
# Used by both /classlist (per-assignment) and /collation (full semester report)
# Accepts CSV, Excel (.xlsx/.xls), Word (.docx), PDF
# Flexible column names: name/full_name/student_name, matric/matric_number/reg_no, dept/department/faculty

def parse_classlist_bytes(raw: bytes, filename: str) -> list[dict]:
    """
    Parse a class list file (CSV / Excel / Word / PDF) and return a list of:
      [{ "full_name": str, "matric_number": str, "department": str }, ...]
    Raises HTTPException with a clear message if anything goes wrong.
    """
    filename = filename.lower()

    NAME_ALIASES   = {"full_name", "name", "student_name", "student name", "fullname"}
    MATRIC_ALIASES = {"matric_number", "matric", "reg_no", "reg number", "registration number",
                      "regno", "matric no", "matricnumber"}
    DEPT_ALIASES   = {"department", "dept", "faculty", "department name"}

    def normalize_header(headers):
        mapping = {}
        for i, h in enumerate(headers):
            hl = h.strip().lower()
            if hl in NAME_ALIASES:
                mapping["full_name"] = i
            elif hl in MATRIC_ALIASES:
                mapping["matric_number"] = i
            elif hl in DEPT_ALIASES:
                mapping["department"] = i
        return mapping

    def validate_mapping(mapping):
        missing_cols = [k for k in ["full_name", "matric_number", "department"] if k not in mapping]
        if missing_cols:
            readable = {
                "full_name": "name/full_name",
                "matric_number": "matric/matric_number/reg_no",
                "department": "department/dept/faculty"
            }
            raise HTTPException(
                status_code=400,
                detail=f"Could not find these columns: {', '.join(readable[c] for c in missing_cols)}. "
                       f"Please make sure your file has columns for name, matric number, and department."
            )

    students = []

    try:
        if filename.endswith(".csv"):
            lines = raw.decode("utf-8-sig").strip().splitlines()
            if not lines:
                raise HTTPException(status_code=400, detail="File is empty")
            headers = [h.strip() for h in lines[0].split(",")]
            mapping = normalize_header(headers)
            validate_mapping(mapping)
            for line in lines[1:]:
                if not line.strip():
                    continue
                parts = line.split(",")
                entry = {
                    "full_name": parts[mapping["full_name"]].strip() if mapping["full_name"] < len(parts) else "",
                    "matric_number": parts[mapping["matric_number"]].strip() if mapping["matric_number"] < len(parts) else "",
                    "department": parts[mapping["department"]].strip() if mapping["department"] < len(parts) else "",
                }
                if entry["matric_number"]:
                    students.append(entry)

        elif filename.endswith(".xlsx") or filename.endswith(".xls"):
            wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                raise HTTPException(status_code=400, detail="Excel file is empty")
            headers = [str(c).strip() if c else "" for c in rows[0]]
            mapping = normalize_header(headers)
            validate_mapping(mapping)
            for row in rows[1:]:
                vals = [str(c).strip() if c is not None else "" for c in row]
                if not any(vals):
                    continue
                entry = {
                    "full_name": vals[mapping["full_name"]] if mapping["full_name"] < len(vals) else "",
                    "matric_number": vals[mapping["matric_number"]] if mapping["matric_number"] < len(vals) else "",
                    "department": vals[mapping["department"]] if mapping["department"] < len(vals) else "",
                }
                if entry["matric_number"]:
                    students.append(entry)

        elif filename.endswith(".docx"):
            doc = DocxDocument(io.BytesIO(raw))
            if not doc.tables:
                raise HTTPException(status_code=400, detail="Word document must contain a table with student data")
            table = doc.tables[0]
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            mapping = normalize_header(headers)
            validate_mapping(mapping)
            for row in table.rows[1:]:
                vals = [cell.text.strip() for cell in row.cells]
                entry = {
                    "full_name": vals[mapping["full_name"]] if mapping["full_name"] < len(vals) else "",
                    "matric_number": vals[mapping["matric_number"]] if mapping["matric_number"] < len(vals) else "",
                    "department": vals[mapping["department"]] if mapping["department"] < len(vals) else "",
                }
                if entry["matric_number"]:
                    students.append(entry)

        elif filename.endswith(".pdf"):
            rows_found = []
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                for page in pdf.pages:
                    for table in (page.extract_tables() or []):
                        for row in table:
                            rows_found.append([str(c).strip() if c else "" for c in row])
            if not rows_found:
                raise HTTPException(status_code=400, detail="Could not find a table in the PDF.")
            headers = rows_found[0]
            mapping = normalize_header(headers)
            validate_mapping(mapping)
            for row in rows_found[1:]:
                entry = {
                    "full_name": row[mapping["full_name"]] if mapping["full_name"] < len(row) else "",
                    "matric_number": row[mapping["matric_number"]] if mapping["matric_number"] < len(row) else "",
                    "department": row[mapping["department"]] if mapping["department"] < len(row) else "",
                }
                if entry["matric_number"]:
                    students.append(entry)

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file. Please upload CSV, Excel (.xlsx), Word (.docx), or PDF."
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read file: {str(e)}")

    if not students:
        raise HTTPException(status_code=400, detail="No student records found in the file")

    return students


@app.post("/assignments/{assignment_id}/classlist")
async def upload_classlist(
    assignment_id: str,
    file: UploadFile = File(...),
    user=Depends(get_current_user)
):
    """
    Per-assignment class list check.
    Upload the full class list → get back who submitted and who is missing for THIS assignment.
    Accepts CSV, Excel, Word, PDF with flexible column names.
    """
    raw = await file.read()
    students = parse_classlist_bytes(raw, file.filename or "")

    subs = supabase.table("submissions").select("matric_number") \
        .eq("assignment_id", assignment_id).execute().data
    submitted_matrics = {s["matric_number"].upper() for s in subs}

    missing = [s for s in students if s.get("matric_number", "").upper() not in submitted_matrics]

    return {
        "total": len(students),
        "submitted": len(students) - len(missing),
        "missing": missing
    }


@app.get("/files/{submission_id}")
async def file_proxy(
    submission_id: str,
    token: str = None,
    credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer(auto_error=False))
):
    """
    Proxy file view — called by frontend as /files/{submission_id}?token=xxx
    Hides raw Supabase storage URL completely.
    Accepts token via query param (for window.open) or Authorization header.
    """
    # Accept token from query param or Authorization header
    raw_token = token or (credentials.credentials if credentials else None)
    if not raw_token:
        raise HTTPException(status_code=401, detail="Authentication required")

    # Validate token (custom JWT or Supabase JWT)
    try:
        payload = jwt.decode(raw_token, SECRET_KEY, algorithms=[ALGORITHM])
        if not payload.get("sub"):
            raise ValueError("no sub")
    except Exception:
        try:
            user = supabase.auth.get_user(raw_token)
            if not user or not user.user:
                raise HTTPException(status_code=401, detail="Invalid token")
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=401, detail="Invalid token")

    sub = supabase.table("submissions").select("file_url, matric_number, full_name")\
        .eq("id", submission_id).single().execute()
    if not sub.data:
        raise HTTPException(status_code=404, detail="Submission not found")

    file_url = sub.data["file_url"]
    async with httpx.AsyncClient() as client:
        resp = await client.get(file_url, timeout=30)

    if resp.status_code != 200:
        raise HTTPException(status_code=404, detail="File not found in storage")

    ext = file_url.split(".")[-1].split("?")[0].lower()
    content_type_map = {
        "pdf": "application/pdf",
        "doc": "application/msword",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }
    content_type = content_type_map.get(ext, "application/octet-stream")
    safe_name = re.sub(r"[^a-zA-Z0-9._-]", "_", f"{sub.data['matric_number']}_{sub.data['full_name']}.{ext}")

    return StreamingResponse(
        io.BytesIO(resp.content),
        media_type=content_type,
        headers={"Content-Disposition": f"inline; filename={safe_name}"}
    )


@app.get("/export/{assignment_id}/excel")
async def export_excel(
    assignment_id: str,
    department: Optional[str] = None,
    user=Depends(get_current_user)
):
    assignment = supabase.table("assignments").select("*").eq("id", assignment_id).single().execute()
    if not assignment.data:
        raise HTTPException(status_code=404, detail="Assignment not found")
    a = assignment.data

    query = supabase.table("submissions").select("*").eq("assignment_id", assignment_id)
    if department:
        query = query.eq("department", department)
    subs = query.order("department").order("group_number").order("full_name").execute().data

    if not subs:
        raise HTTPException(status_code=404, detail="No submissions found")

    def make_excel_sheet(rows, dept_name):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = dept_name[:31]

        ws.merge_cells("A1:I1")
        ws["A1"] = f"{a['course_name']} — {a['title']}"
        ws["A1"].font = Font(bold=True, size=13)
        ws.merge_cells("A2:I2")
        ws["A2"] = f"Department: {dept_name}  |  Deadline: {a['deadline'][:16].replace('T', ' ')}"
        ws["A2"].font = Font(size=9, italic=True)

        headers = ["S/N", "Full Name", "Matric Number", "Department", "Group", "Submitted At", "Late?", "Score", "File Link"]
        hrow = 4
        hfill = PatternFill("solid", fgColor="1A3A5C")
        hfont = Font(bold=True, color="FFFFFF")
        thin = Side(style="thin", color="DDDDDD")
        bdr = Border(left=thin, right=thin, top=thin, bottom=thin)

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=hrow, column=col, value=h)
            cell.fill = hfill
            cell.font = hfont
            cell.alignment = Alignment(horizontal="center")
            cell.border = bdr

        alt = PatternFill("solid", fgColor="EBF2FA")
        for i, row in enumerate(rows, 1):
            r = hrow + i
            fill = alt if i % 2 == 0 else PatternFill()
            vals = [
                i,
                row.get("full_name", ""),
                row.get("matric_number", ""),
                row.get("department", ""),
                f"Group {row['group_number']}" if row.get("group_number") else "Individual",
                str(row.get("submitted_at", ""))[:16].replace("T", " "),
                "YES" if row.get("is_late") else "No",
                row.get("score", ""),
                row.get("file_url", ""),
            ]
            for col, val in enumerate(vals, 1):
                cell = ws.cell(row=r, column=col, value=val)
                cell.fill = fill
                cell.border = bdr
                if col == 9:
                    matric = row.get("matric_number", "")
                    file_url = row.get("file_url", "")
                    if file_url:
                        # Write actual clickable hyperlink into the cell
                        cell.value = f"View File - {matric}"
                        cell.hyperlink = file_url
                        cell.font = Font(color="0563C1", underline="single")
                    else:
                        cell.value = "No file"

        widths = [5, 26, 18, 30, 12, 20, 7, 8, 55]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf

    if department:
        buf = make_excel_sheet(subs, department)
        safe_dept = re.sub(r'[^a-zA-Z0-9]', '_', department)
        filename = f"{safe_dept}_{a['course_name']}.xlsx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    else:
        by_dept = defaultdict(list)
        for s in subs:
            by_dept[s["department"]].append(s)

        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            for dept, rows in by_dept.items():
                excel_buf = make_excel_sheet(rows, dept)
                safe_dept = re.sub(r'[^a-zA-Z0-9]', '_', dept)
                zf.writestr(f"{safe_dept}.xlsx", excel_buf.read())

        zip_buf.seek(0)
        safe_title = re.sub(r'[^a-zA-Z0-9]', '_', a["title"])
        return StreamingResponse(
            zip_buf,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={safe_title}_Excel_Export.zip"}
        )


# ─── ADMIN ROUTES ─────────────────────────────────────────────────────────────

@app.post("/admin/login")
@limiter.limit("5/minute")
async def admin_login(
    request: Request,
    email: str = Form(...),
    password: str = Form(...),
):
    if email != ADMIN_EMAIL or password != ADMIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid admin credentials")

    timestamp = int(datetime.now(timezone.utc).timestamp())
    signature = hashlib.sha256(f"admin:{timestamp}:{ADMIN_SECRET}".encode()).hexdigest()[:32]
    token = f"admin:{timestamp}:{signature}"

    return {"token": token, "message": "Admin login successful"}


@app.get("/admin/lecturers")
async def admin_get_lecturers(is_admin=Depends(verify_admin_token)):
    try:
        # list_users with per_page to ensure we get all users
        # Supabase SDK v2+ returns paginated object with .users attribute
        try:
            users_response = supabase.auth.admin.list_users(page=1, per_page=1000)
        except TypeError:
            users_response = supabase.auth.admin.list_users()

        if hasattr(users_response, "users"):
            users_list = users_response.users
        elif isinstance(users_response, list):
            users_list = users_response
        else:
            try:
                users_list = list(users_response)
            except Exception:
                users_list = []

        lecturers = []

        for u in users_list:
            if not u.email:
                continue

            assignment_result = supabase.table("assignments")\
                .select("id", count="exact")\
                .eq("lecturer_id", str(u.id))\
                .execute()
            assignment_count = assignment_result.count or 0

            submission_count = 0
            if assignment_count > 0:
                assignments = supabase.table("assignments")\
                    .select("id")\
                    .eq("lecturer_id", str(u.id))\
                    .execute().data
                assignment_ids = [a["id"] for a in assignments]
                for aid in assignment_ids:
                    sub_result = supabase.table("submissions")\
                        .select("id", count="exact")\
                        .eq("assignment_id", aid)\
                        .execute()
                    submission_count += sub_result.count or 0

            full_name = u.user_metadata.get("full_name", "Unknown") if u.user_metadata else "Unknown"
            lecturers.append({
                "id": str(u.id),
                "email": u.email,
                "full_name": full_name,
                "display_name": full_name,
                "is_verified": u.email_confirmed_at is not None,
                "created_at": u.created_at.isoformat() if u.created_at else "",
                "assignment_count": assignment_count,
                "submission_count": submission_count,
            })

        lecturers.sort(key=lambda x: x["created_at"], reverse=True)
        return lecturers

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch lecturers: {str(e)}")


@app.post("/admin/impersonate/{lecturer_id}")
async def admin_impersonate(
    lecturer_id: str,
    is_admin=Depends(verify_admin_token)
):
    try:
        user_response = supabase.auth.admin.get_user_by_id(lecturer_id)

        if not user_response or not user_response.user:
            raise HTTPException(status_code=404, detail="Lecturer not found")

        user = user_response.user

        token = create_access_token({
            "sub": str(user.id),
            "email": user.email,
            "role": "lecturer",
            "impersonated": True
        })

        return {"token": token}

    except HTTPException:
        raise
    except Exception as e:
        print("ERROR:", str(e))
        raise HTTPException(status_code=500, detail="Impersonation failed")


# ─── Collation ────────────────────────────────────────────────────────────────

@app.post("/collation")
async def collate_assignments(
    assignment_ids: str = Form(...),   # comma-separated assignment IDs
    format: str = Form("excel"),       # excel or csv
    filter_department: Optional[str] = Form(None),  # filter by department
    filter_level: Optional[str] = Form(None),       # filter by level e.g. "100L", "200L"
    search_matric: Optional[str] = Form(None),      # search one student by matric
    classlist_file: Optional[UploadFile] = File(None),  # optional: full class register
    user=Depends(get_current_user)
):
    """
    Collation report across multiple assignments.
    - filter_department: partial match on student's submitted department
    - filter_level: matches against assignment's target_level tag (set at creation)
      e.g. filtering "100L" will only include assignments tagged as 100L
    - search_matric: search one student across all selected assignments
    - classlist_file: optional class register (CSV/Excel/Word/PDF).
      When provided, EVERY student in the register appears in the report —
      even those who submitted nothing (they get all NO ❌).
      Without it, only students who submitted at least once appear.
    """
    ids = [a.strip() for a in assignment_ids.split(",") if a.strip()]
    if not ids:
        raise HTTPException(status_code=400, detail="No assignment IDs provided")

    # Fetch assignment details — only assignments belonging to this lecturer
    # Now includes target_level so we can filter properly
    assignments = []
    for aid in ids:
        res = supabase.table("assignments") \
            .select("id, title, course_name, target_level") \
            .eq("id", aid) \
            .eq("lecturer_id", str(user.id)) \
            .single().execute()
        if res.data:
            assignments.append(res.data)

    if not assignments:
        raise HTTPException(status_code=404, detail="No valid assignments found")

    # ── Filter assignments by level BEFORE fetching submissions ───────────────
    # This is the correct approach: filter which assignments count toward the report
    # rather than guessing level from matric numbers (unreliable)
    if filter_level and filter_level.strip():
        lvl = filter_level.strip().upper()
        # Normalize: "100" → "100L", "100l" → "100L"
        if lvl.isdigit():
            lvl = lvl + "L"
        lvl_num = lvl.rstrip("L")  # "100" from "100L"

        filtered_assignments = []
        for a in assignments:
            al = (a.get("target_level") or "").strip().upper()
            # Match if assignment level equals filter OR if no level tagged (include all untagged)
            if not al:
                # Untagged assignments — include them (don't exclude just because lecturer forgot to tag)
                filtered_assignments.append(a)
            elif al == lvl or al == lvl_num or lvl_num in al:
                filtered_assignments.append(a)

        if not filtered_assignments:
            raise HTTPException(
                status_code=404,
                detail=f"None of the selected assignments are tagged for Level {filter_level}. "
                       f"Tag assignments with their target level when creating them, or remove the level filter."
            )
        assignments = filtered_assignments

    # Fetch all submissions for selected assignments
    # student_map: matric -> {full_name, matric_number, department, assignment_id: True/False}
    student_map = {}

    for a in assignments:
        subs = supabase.table("submissions") \
            .select("matric_number, full_name, department") \
            .eq("assignment_id", a["id"]).execute().data

        for s in subs:
            m = s["matric_number"].upper()
            if m not in student_map:
                student_map[m] = {
                    "full_name": s["full_name"],
                    "matric_number": m,
                    "department": s["department"],
                }
            student_map[m][a["id"]] = True

    # ── Merge class list if uploaded ──────────────────────────────────────────
    # This is the key upgrade: every student in the register appears in the report,
    # even if they submitted NOTHING. Without a class list, zero-submission students
    # are invisible. With it, they show as all NO ❌ — the lecturer sees who did nothing.
    classlist_students = []
    if classlist_file and classlist_file.filename:
        raw_cl = await classlist_file.read()
        classlist_students = parse_classlist_bytes(raw_cl, classlist_file.filename)
        for s in classlist_students:
            m = s["matric_number"].upper()
            if m not in student_map:
                # Student has zero submissions — add them with empty assignment slots
                student_map[m] = {
                    "full_name": s["full_name"],
                    "matric_number": m,
                    "department": s["department"],
                    "_from_classlist": True,  # flag so we can style them differently if needed
                }
            # If student IS in student_map (submitted something), update name/dept
            # from class list in case they typed their name slightly differently
            else:
                # Trust the class list for name and department (lecturer's official record)
                student_map[m]["full_name"] = s["full_name"]
                student_map[m]["department"] = s["department"]

    # ── Filter by department (case-insensitive partial match on submitted dept) ─
    if filter_department and filter_department.strip():
        dept_lower = filter_department.strip().lower()
        student_map = {
            m: info for m, info in student_map.items()
            if dept_lower in info.get("department", "").lower()
        }

    # ── Filter by specific student matric number ──────────────────────────────
    if search_matric and search_matric.strip():
        sm = search_matric.strip().upper()
        student_map = {
            m: info for m, info in student_map.items()
            if sm in m
        }

    if not student_map:
        raise HTTPException(
            status_code=404,
            detail="No students found matching your filters for the selected assignments"
        )

    # ── Build report rows ─────────────────────────────────────────────────────
    rows = []
    for matric, info in student_map.items():
        row = {
            "Full Name": info["full_name"],
            "Matric Number": info["matric_number"],
            "Department": info["department"],
        }
        total_done = 0
        for a in assignments:
            did_it = info.get(a["id"], False)
            level_tag = f" [{a['target_level']}]" if a.get("target_level") else ""
            col_name = f"{a['course_name']} - {a['title']}{level_tag}"
            row[col_name] = "YES ✅" if did_it else "NO ❌"
            if did_it:
                total_done += 1
        row["Submitted"] = total_done
        row["Missed"] = len(assignments) - total_done
        row["Completion"] = f"{total_done}/{len(assignments)}"
        rows.append(row)

    # Sort: department first, then most missed, then name
    rows.sort(key=lambda x: (x["Department"], -x["Missed"], x["Full Name"]))

    headers = list(rows[0].keys())
    assignment_cols = [
        f"{a['course_name']} - {a['title']}" + (f" [{a['target_level']}]" if a.get("target_level") else "")
        for a in assignments
    ]

    # ── Excel output ──────────────────────────────────────────────────────────
    if format == "excel":
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Collation Report"

        # Title
        title_col = openpyxl.utils.get_column_letter(len(headers))
        ws.merge_cells(f"A1:{title_col}1")
        ws["A1"] = "Assignment Collation Report"
        ws["A1"].font = Font(bold=True, size=14)
        ws["A1"].alignment = Alignment(horizontal="center")

        # Subtitle showing filters applied
        subtitle_parts = []
        if filter_department:
            subtitle_parts.append(f"Department: {filter_department}")
        if filter_level:
            subtitle_parts.append(f"Level: {filter_level}")
        if search_matric:
            subtitle_parts.append(f"Student: {search_matric}")
        subtitle_parts.append(f"Assignments: {len(assignments)} selected")
        subtitle_parts.append(f"Total Students: {len(rows)}")
        if classlist_students:
            subtitle_parts.append(f"Class Register: {len(classlist_students)} students (includes non-submitters)")

        ws.merge_cells(f"A2:{title_col}2")
        ws["A2"] = "  |  ".join(subtitle_parts)
        ws["A2"].font = Font(size=9, italic=True)

        # Assignment list row
        ws.merge_cells(f"A3:{title_col}3")
        ws["A3"] = "  |  ".join([f"{a['course_name']}: {a['title']}" for a in assignments])
        ws["A3"].font = Font(size=8, italic=True, color="555555")

        # Headers row (row 5)
        hfill = PatternFill("solid", fgColor="1A3A5C")
        hfont = Font(bold=True, color="FFFFFF")
        thin = Side(style="thin", color="DDDDDD")
        bdr = Border(left=thin, right=thin, top=thin, bottom=thin)

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=h)
            cell.fill = hfill
            cell.font = hfont
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
            cell.border = bdr

        # Data rows — group by department with dept header rows
        green_fill = PatternFill("solid", fgColor="C6EFCE")
        red_fill   = PatternFill("solid", fgColor="FFC7CE")
        alt_fill   = PatternFill("solid", fgColor="EBF2FA")
        dept_fill  = PatternFill("solid", fgColor="2D5F4E")

        current_dept = None
        data_row = 6

        for i, row in enumerate(rows):
            # Insert department header when dept changes
            if row["Department"] != current_dept:
                current_dept = row["Department"]
                ws.merge_cells(f"A{data_row}:{title_col}{data_row}")
                dept_cell = ws.cell(row=data_row, column=1, value=f"  {current_dept.upper()}")
                dept_cell.fill = dept_fill
                dept_cell.font = Font(bold=True, color="FFFFFF", size=10)
                dept_cell.alignment = Alignment(vertical="center")
                ws.row_dimensions[data_row].height = 20
                data_row += 1

            fill = alt_fill if i % 2 == 0 else PatternFill()
            for col, h in enumerate(headers, 1):
                val = row[h]
                cell = ws.cell(row=data_row, column=col, value=val)
                cell.border = bdr
                cell.alignment = Alignment(horizontal="center", vertical="center")
                if h in assignment_cols:
                    if str(val).startswith("YES"):
                        cell.fill = green_fill
                        cell.font = Font(color="276221", bold=True)
                    else:
                        cell.fill = red_fill
                        cell.font = Font(color="9C0006", bold=True)
                elif h == "Completion":
                    cell.font = Font(bold=True)
                    cell.fill = fill
                else:
                    cell.fill = fill

            data_row += 1

        # Column widths
        col_widths = {"Full Name": 25, "Matric Number": 16, "Department": 28,
                      "Submitted": 10, "Missed": 8, "Completion": 12}
        for col, h in enumerate(headers, 1):
            width = col_widths.get(h, 20)
            ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width

        ws.freeze_panes = "A6"  # freeze header rows

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=collation_report.xlsx"}
        )

    # ── CSV output ────────────────────────────────────────────────────────────
    elif format == "csv":
        output = io.StringIO()
        output.write(",".join(f'"{h}"' for h in headers) + "")
        current_dept = None
        for row in rows:
            if row["Department"] != current_dept:
                current_dept = row["Department"]
                output.write(f'"--- {current_dept} ---"')
            output.write(",".join([f'"{str(row[h])}"' for h in headers]) + "")
        output.seek(0)
        return StreamingResponse(
            io.BytesIO(output.getvalue().encode("utf-8-sig")),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename=collation_report.csv"}
        )

    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use excel or csv.")


# ─── Health check ─────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "app": "Assignify API"}